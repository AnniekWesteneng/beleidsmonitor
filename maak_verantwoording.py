"""Genereer een uitgebreid verantwoordingsdocument (gewoon Word, geen huisstijl)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ZWART = RGBColor(0, 0, 0)
GRIJS = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "DDDDDD"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def kop(doc, tekst, niveau=1):
    h = doc.add_heading(tekst, level=niveau)
    for r in h.runs:
        r.font.color.rgb = ZWART
    return h


def para(doc, tekst):
    return doc.add_paragraph(tekst)


def bullet(doc, items):
    for it in items:
        if isinstance(it, tuple):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(it[0] + " — ").bold = True
            p.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")


def tabel(doc, headers, rijen, breedtes):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(10)
        shade(c, HEADER_FILL)
    for rij in rijen:
        cells = t.add_row().cells
        for i, w in enumerate(rij):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(w); run.font.size = Pt(10)
    for row in t.rows:
        for i, b in enumerate(breedtes):
            row.cells[i].width = Cm(b)
    return t


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.5)
sec.top_margin = sec.bottom_margin = Cm(2.5)
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)

# Titel
tp = doc.add_paragraph(); tr = tp.add_run("Verantwoordingsdocument — Bouw van de Beleidsmonitor")
tr.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = ZWART
st = doc.add_paragraph(); s = st.add_run("Industrieel vastgoed · Today Development · ontwikkeling juni 2026")
s.italic = True; s.font.color.rgb = GRIJS

# 1. Doel
kop(doc, "1. Doel van dit document")
para(doc, "Dit document legt vast hoe de beleidsmonitor is gebouwd en, belangrijker, "
          "waarom voor elke stap bepaalde keuzes zijn gemaakt. Het beschrijft ook expliciet "
          "welke onderdelen tijdens de ontwikkeling zijn ontstaan en dus afweken van of "
          "verder gingen dan het oorspronkelijke bouwplan. Zo is het hele bouwproces "
          "navolgbaar en verantwoord.")

# 2. Uitgangspunt
kop(doc, "2. Uitgangspunt: het oorspronkelijke bouwplan")
para(doc, "Het project startte met een vooraf opgesteld bouwplan in zeven fasen: "
          "(0) projectopzet, (1) configuratie met 3 gemeenten en 10 indicatoren, "
          "(2) eerste bron Officiële Bekendmakingen, (3) AI-classificatie, (4) SQLite-database, "
          "(5) pipeline, (6) Streamlit-dashboard, en (7) opschalen (meer gemeenten, een tweede "
          "bron, automatisering). Het leidende principe was: klein beginnen en stap voor stap "
          "testen vóór uitbreiden. Tijdens de uitvoering zijn op grond van tussentijdse "
          "bevindingen onderdelen toegevoegd of aangescherpt; die staan in hoofdstuk 6.")

# 3. Hulpmiddelen en werkwijze
kop(doc, "3. Hulpmiddelen en werkwijze")
para(doc, "De programmacode is geschreven met Claude Code (Anthropic), een AI-assistent voor "
          "softwareontwikkeling, op basis van mijn specificaties en aansturing. Ik heb het "
          "inhoudelijke kader (de 10 indicatoren), de bron- en scopekeuzes en alle "
          "ontwerpbeslissingen bepaald; de AI vertaalde dit naar code en deed suggesties.")
para(doc, "Een centrale werkwijze was bron-verificatie vóór bouw: bij elke databron is eerst "
          "met een testscript de echte API-respons opgehaald en bekeken, zodat de verwerking op "
          "de werkelijke structuur is gebaseerd en niet op aannames. Deze verkenningsscripts "
          "(bronnen/_verken_*.py) zijn bewaard als bewijs.")
para(doc, "Verantwoording: deze werkwijze voorkomt fouten door verkeerde aannames over "
          "veldnamen of datastructuren, en maakt de betrouwbaarheid van de koppelingen "
          "controleerbaar.")

# 4. Techniekkeuzes
kop(doc, "4. Techniekkeuzes en verantwoording")
tabel(doc, ["Keuze", "Verantwoording", "Overwogen alternatief"],
      [
          ["Python 3.12", "Breed gebruikt, veel open-data-bibliotheken, makkelijk leesbaar.", "—"],
          ["SQLite", "Eén bestand, geen server; ideaal om klein te starten en mee te nemen.", "PostgreSQL (pas nodig bij groot volume)"],
          ["Streamlit", "Dashboard met één commando, zonder front-endkennis.", "Eigen webframework (veel meer werk)"],
          ["requests", "Eenvoudige, standaard HTTP-aanroepen.", "—"],
          ["anthropic (SDK)", "Officiële, stabiele koppeling met het Claude-model.", "—"],
          ["bs4 / lxml / pypdf", "Tekst uit XML, HTML en PDF kunnen halen.", "—"],
      ],
      [3.5, 8.5, 4.0])

# 5. Bouw per onderdeel
kop(doc, "5. Bouw per onderdeel, met verantwoording")

kop(doc, "5.1 Configuratie en indicatoren", 2)
para(doc, "Alle variabele instellingen (gemeenten, provincies, indicatoren, zoektermen, "
          "vanaf-jaar) staan centraal in config.py. Verantwoording: opschalen of bijsturen "
          "betekent dan alleen een lijst aanpassen, zonder de rest van de code te raken. De 10 "
          "indicatoren komen uit het onderzoekskader; ze zijn ongewijzigd overgenomen.")

kop(doc, "5.2 Bron 1 — Officiële Bekendmakingen (SRU)", 2)
para(doc, "Als eerste bron gekozen omdat het de best gestructureerde, gratis en dagelijks "
          "geactualiseerde bron is (XML via de SRU-API). Geverifieerde keuzes:")
bullet(doc, [
    ("Filteren op gemeente via dt.creator", "uit de verkenning bleek dit het werkende veld; exacte match (==) gaf geen resultaten."),
    ("Datumfilter (dt.date >= vanaf-jaar) + nieuwste eerst", "oudere stukken (bv. 2010) zijn voor een actuele kans/risico-inschatting niet relevant; gebruiker gaf dit expliciet aan."),
    ("Provinciaal blad onderscheiden", "gemeente en provincie delen soms een naam (Utrecht); zonder onderscheid zouden provinciale stukken de gemeentelijke vervuilen."),
    ("Den Haag onder twee namen", "uit de data bleek Den Haag zowel als 'Den Haag' als ''s-Gravenhage' voor te komen; beide worden opgehaald."),
    ("Retry-met-backoff", "de dienst throttelt bij snelle bevraging; met pauzes en herhaalpogingen blijft het ophalen betrouwbaar."),
])

kop(doc, "5.3 Tekstextractie", 2)
para(doc, "De SRU-records bevatten alleen metadata; de volledige tekst wordt apart opgehaald in "
          "de volgorde XML, dan HTML, dan PDF (eerste die lukt). Verantwoording: XML is het "
          "schoonst; PDF als laatste vangnet zodat ook alleen-PDF-documenten meetellen. Voor "
          "lange documenten wordt niet bot afgekapt maar een fragment rond de zoekterm gekozen, "
          "zodat relevante passages verderop niet wegvallen.")

kop(doc, "5.4 Classificatie", 2)
para(doc, "Elk document wordt door het Claude-model getoetst aan de 10 indicatoren via een "
          "vaste systeeminstructie met de definities van kans, risico en contextafhankelijk. "
          "Belangrijke keuzes en verantwoording:")
bullet(doc, [
    ("Model claude-sonnet-4-6", "het in het oorspronkelijke plan genoemde model (claude-sonnet-4-20250514) bleek per 2026 uitgefaseerd; via de API is het beschikbare, actuele opvolgermodel gekozen."),
    ("Antwoord uitsluitend in JSON", "maakt de uitvoer machineleesbaar en betrouwbaar te verwerken; parsefouten worden netjes afgevangen."),
    ("Meerdere indicatoren per document", "afwijking van het oorspronkelijke plan (dat één indicator per document gaf): in de praktijk raakt één document vaak meerdere indicatoren, dus dit verhoogt de volledigheid."),
    ("Relevantiescore 1-5", "toegevoegd om de signalen te kunnen rangschikken op belang; niet in het oorspronkelijke plan."),
    ("Tekstlengte begrensd", "om kosten te beheersen; instelbaar via MAX_TEKST_TEKENS."),
])

kop(doc, "5.5 Database", 2)
para(doc, "SQLite met één tabel 'signalen'. Omdat één document meerdere indicator-signalen kan "
          "opleveren, ligt de uniciteit op de combinatie (url, indicator_id) in plaats van op url "
          "alleen. Verantwoording: zo worden dubbele rijen voorkomen én kunnen meerdere "
          "indicatoren per document worden bewaard. Een controle op al verwerkte URLs voorkomt "
          "dat documenten dubbel (en duur) worden geclassificeerd.")

kop(doc, "5.6 Pipeline", 2)
para(doc, "De pipeline koppelt de lagen: per bron, per gebied (gemeente, en provincie waar de "
          "bron dat ondersteunt), per zoekterm worden documenten opgehaald, al verwerkte "
          "overgeslagen, nieuwe geclassificeerd en relevante opgeslagen. Verantwoording: deze "
          "opzet is bron-agnostisch; een nieuwe bron toevoegen vergt alleen een module die "
          "dezelfde interface biedt.")

kop(doc, "5.7 Dashboard", 2)
para(doc, "Streamlit-dashboard met drie tabbladen (Signalen, Netcongestie, Chat). Keuzes:")
bullet(doc, [
    ("Filters in een zijbalk", "meer ruimte voor de signalen; consistent filter over de tabbladen."),
    ("Vrije-tekst zoeken, datumbereik, sorteren", "toegevoegd voor bruikbaarheid bij groeiende datasets."),
    ("Relevantiefilter standaard op 4", "toont meteen alleen de sterke signalen en onderdrukt ruis."),
    ("Signalen groeperen per document", "voorkomt een lange lijst losse regels nu één document meerdere indicatoren geeft."),
    ("Indicatornaam tonen i.p.v. nummer", "leesbaarder dan 'indicator 2'."),
    ("Wachtwoord-slot en secrets", "nodig om het dashboard veilig te kunnen delen; de API-sleutel staat nooit in de code."),
])
para(doc, "Toegevoegd buiten het oorspronkelijke plan: de huisstijl (kleuren en logo-woordmerk "
          "van Today Development) en de chatfunctie waarmee in natuurlijke taal vragen over de "
          "signalen gesteld kunnen worden.")

kop(doc, "5.8 Bron 2 — Open Raadsinformatie", 2)
para(doc, "In het oorspronkelijke plan stond deze tweede bron pas in de opschaalfase (Fase 7); "
          "hij is eerder toegevoegd omdat hij beleid opvangt dat vaak vóór officiële publicatie "
          "al in de gemeenteraad speelt. De API (Elasticsearch) is eerst geverifieerd: per "
          "gemeente een eigen index, volledige tekst in het veld 'text'. Verantwoording: de twee "
          "bronnen vullen elkaar aan — gepubliceerde besluiten versus wat nog in de raad speelt.")

kop(doc, "5.9 Bron 3 — Netcongestie (ontstaan tijdens de ontwikkeling)", 2)
para(doc, "Deze bron stond NIET in het oorspronkelijke bouwplan. Tijdens de ontwikkeling kwam "
          "dit idee naar voren (voorgesteld door de AI-assistent tijdens een verkenning van "
          "aanvullende bronnen) toen bleek dat indicator 4 (energie-infrastructuur en "
          "netcongestie) inhoudelijk grotendeels buiten gemeentelijk beleid valt en daardoor "
          "ondervertegenwoordigd was. Na akkoord is de capaciteitskaart van Netbeheer Nederland "
          "toegevoegd.")
para(doc, "Belangrijke keuze met verantwoording: netcongestie is geen documentenbron maar live "
          "statusdata (een open ArcGIS-kaart). Daarom is het NIET door de AI-classificatie "
          "gehaald, maar als apart statuspaneel getoond. De koppeling met de gemeente gebeurt op "
          "naam van het voedingsgebied — een bewuste vereenvoudiging; een exacte geografische "
          "koppeling is als verbeterpunt voor later vastgelegd. Bij het paneel staat een "
          "disclaimer en een link naar de officiële kaart, zodat geen valse precisie ontstaat.")

# 6. Afwijkingen
kop(doc, "6. Afwijkingen van en uitbreidingen op het oorspronkelijke plan")
para(doc, "Onderstaande onderdelen zijn tijdens de ontwikkeling toegevoegd of aangepast op "
          "grond van tussentijdse bevindingen. Per onderdeel de reden.")
tabel(doc, ["Onderdeel", "In origineel plan?", "Reden / verantwoording"],
      [
          ["Netcongestie-paneel", "Nee (ontstaan onderweg)", "Indicator 4 bleek ondervertegenwoordigd; live netstatus toegevoegd als apart paneel."],
          ["Provinciale bron", "Nee", "Provincies programmeren veel bedrijventerreinen; relevant beleid werd anders gemist."],
          ["Meerdere indicatoren per document", "Nee (plan: één)", "Een document raakt vaak meerdere indicatoren; verhoogt volledigheid."],
          ["Relevantiescore + sorteren", "Nee", "Om signalen op belang te kunnen rangschikken en ruis te onderdrukken."],
          ["Uitgebreide zoektermen + synoniemen", "Deels (plan: 6 termen)", "Bredere dekking; de zoekstap (trefwoorden) bleek de grootste beperking voor volledigheid."],
          ["PDF-extractie en slim tekstfragment", "Nee", "Anders vielen alleen-PDF-documenten en lange documenten (deels) weg."],
          ["Dashboard-uitbreidingen", "Deels", "Filters, zoeken, datumbereik, sorteren, groeperen per document — voor bruikbaarheid."],
          ["Chatfunctie", "Nee", "Vragen in natuurlijke taal over de signalen kunnen stellen."],
          ["Huisstijl", "Nee", "Presentabel maken voor het bedrijf en de scriptie."],
          ["Online via Streamlit Cloud + wachtwoord", "Deels (plan: GitHub Actions)", "Om het dashboard veilig met collega's te kunnen delen."],
          ["Model bijgewerkt", "Aangepast", "Het geplande model was uitgefaseerd; vervangen door het actuele opvolgermodel."],
          ["Tweede bron eerder gebouwd", "Verschoven", "Open Raadsinformatie stond in Fase 7; eerder toegevoegd vanwege de aanvullende waarde."],
      ],
      [4.5, 3.3, 8.2])

# 7. Geparkeerd
kop(doc, "7. Geparkeerde ideeën (voor later)")
bullet(doc, [
    "IBIS Bedrijventerreinen als bron (bronverkenning gedaan; landelijke open data, maar gedateerd/shapefile).",
    "Locatiekaart met kans/risico-symbolen per plek binnen de gemeente (eerste versie gemaakt; uitbreiden naar alle signalen).",
    "Omgevingsloket-koppeling en zoeken op adres.",
    "Netcongestie geografisch koppelen (i.p.v. op naam).",
    "Automatisch draaien (Fase D) en meldingen bij nieuwe risico's.",
])

# 8. Beperkingen
kop(doc, "8. Beperkingen en betrouwbaarheid")
bullet(doc, [
    "De gebruikte bronnen zijn officieel/gezaghebbend; de onzekerheid zit in de verwerking aan onze kant.",
    "De AI-classificatie is een interpretatie, geen objectieve waarheid; steekproefsgewijs te controleren.",
    "De volledigheid (recall) hangt af van de gebruikte zoektermen — trefwoord-zoeken mist documenten zonder die woorden.",
    "Aantallen weerspiegelen deels documentvolume; gebruik ze als richting, niet als exacte zwaarte.",
    "De netcongestie-gemeentekoppeling is op naam (vereenvoudiging); de officiële kaart is leidend.",
])

# 9. Verantwoording AI-gebruik (volgens APA-richtlijn GenAI als onderzoekstool)
kop(doc, "9. Verantwoording gebruik generatieve AI")
para(doc, "Generatieve AI is in dit onderzoek gebruikt als onderzoekstool voor de bouw van de "
          "softwareprototype. Conform de APA-richtlijn voor generatieve AI worden hieronder de "
          "tool, de input en de output beschreven.")

p = doc.add_paragraph()
p.add_run("Tool. ").bold = True
p.add_run("Voor de ontwikkeling is gebruikgemaakt van ")
p.add_run("Claude Code").italic = True
p.add_run(" (Anthropic, 2026), een generatieve-AI-assistent voor softwareontwikkeling "
          "(aangedreven door het model Claude Opus 4.8). De classificatie binnen de applicatie "
          "gebruikt het model Claude Sonnet 4.6.")

p = doc.add_paragraph()
p.add_run("Input. ").bold = True
p.add_run("Een door mij opgesteld bouwplan met functionele specificaties (het indicatorenkader, "
          "de bronkeuze en de gewenste werking) en iteratieve aansturing en feedback tijdens het "
          "bouwproces in juni 2026. Het bouwplan is opgenomen als bijlage (zie Bijlage A).")

p = doc.add_paragraph()
p.add_run("Output. ").bold = True
p.add_run("De programmacode van de applicatie (Python/Streamlit). Aangezien de output "
          "programmeercode is, is deze beschikbaar gemaakt via een repository in plaats van "
          "integraal opgenomen; de verwijzing staat in Bijlage A. Ik heb de code gecontroleerd en "
          "getest, de databronnen geverifieerd (testscripts tegen de echte respons) en de "
          "AI-classificatie steekproefsgewijs vergeleken met mijn eigen vakinhoudelijke "
          "beoordeling. De beperkingen zijn in hoofdstuk 8 benoemd. Ik begrijp de werking van de "
          "applicatie en kan deze toelichten en verantwoorden.")

kop(doc, "10. Reproductie (kort)")
para(doc, "1) virtuele omgeving aanmaken en packages installeren; 2) API-sleutel in .env zetten; "
          "3) data verzamelen met de pipeline; 4) dashboard tonen met Streamlit. De exacte "
          "commando's en code staan in het technische stappenplan (STAPPENPLAN.md).")

# Bijlage A — AI-verantwoording en bronvermelding
kop(doc, "Bijlage A — AI-verantwoording en bronvermelding")
para(doc, "Onderstaande bronvermeldingen volgen het 6-delige format uit de APA-richtlijn voor "
          "generatieve AI (Ontwikkelaar. (jaar). Naam (Versie) [Generatieve AI]. URL). Pas de "
          "auteursnaam zo nodig aan naar je officiële naam/initialen zoals de opleiding vraagt.")

sub = doc.add_paragraph(); sub.add_run("Gebruikte generatieve-AI-tools").bold = True

b = doc.add_paragraph()
b.add_run("Anthropic. (2026). ")
b.add_run("Claude Code").italic = True
b.add_run(" (Opus 4.8) [Generatieve AI]. https://claude.com/claude-code")

b = doc.add_paragraph()
b.add_run("Anthropic. (2026). ")
b.add_run("Claude").italic = True
b.add_run(" (Sonnet 4.6) [Generatieve AI]. https://www.anthropic.com")

sub = doc.add_paragraph(); sub.add_run("De applicatie als eigen output (programmeercode)").bold = True
b = doc.add_paragraph()
b.add_run("Westeneng, A. (2026). ")
b.add_run("Beleidsmonitor").italic = True
b.add_run(" [Computersoftware]. GitHub. https://github.com/AnniekWesteneng/beleidsmonitor")

para(doc, "Toelichting: de programmacode is de output van de tool. Conform de richtlijn (output "
          "anders dan tekst) is deze beschikbaar via bovenstaande repository met mijn eigen naam "
          "als auteur, in plaats van integraal in dit document opgenomen. Het bouwplan (de input) "
          "kan als aparte bijlage worden toegevoegd.")

uit = "Verantwoordingsdocument Beleidsmonitor.docx"
try:
    doc.save(uit); print("Opgeslagen:", uit)
except PermissionError:
    uit2 = "Verantwoordingsdocument Beleidsmonitor (nieuw).docx"
    doc.save(uit2); print("Origineel open in Word; opgeslagen als:", uit2)
