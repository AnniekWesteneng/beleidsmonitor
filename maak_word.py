"""Genereer een net opgemaakt Word-document van het uitbreidingsplan."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x6E, 0x5E, 0x8C)   # gedempt paars (huisstijl)
DONKER = RGBColor(0x1A, 0x1A, 0x1A)
GRIJS = RGBColor(0x66, 0x66, 0x66)
HEADER_FILL = "6E5E8C"
ZEBRA_FILL = "F2EFF6"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def zet_cel(cell, tekst, vet=False, wit=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(tekst)
    r.bold = vet
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if wit else DONKER


def kop(doc, tekst, niveau=1):
    h = doc.add_heading(tekst, level=niveau)
    for r in h.runs:
        r.font.color.rgb = ACCENT if niveau == 1 else DONKER
    return h


def alinea(doc, tekst, italic=False, grijs=False):
    p = doc.add_paragraph()
    r = p.add_run(tekst)
    r.italic = italic
    r.font.color.rgb = GRIJS if grijs else DONKER
    if grijs:
        r.font.size = Pt(9)
    return p


def bullets(doc, items):
    for it in items:
        if isinstance(it, tuple):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(it[0] + ": ").bold = True
            p.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")


def tabel(doc, headers, rijen, breedtes):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        zet_cel(t.rows[0].cells[i], h, vet=True, wit=True)
        shade(t.rows[0].cells[i], HEADER_FILL)
    for ri, rij in enumerate(rijen):
        cells = t.add_row().cells
        for i, waarde in enumerate(rij):
            zet_cel(cells[i], waarde)
            if ri % 2 == 1:
                shade(cells[i], ZEBRA_FILL)
    for row in t.rows:
        for i, b in enumerate(breedtes):
            row.cells[i].width = Cm(b)
    return t


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.5)
sec.top_margin = sec.bottom_margin = Cm(2.5)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = DONKER

# --- Titel ---
tp = doc.add_paragraph(); tr = tp.add_run("Uitbreidingsplan")
tr.bold = True; tr.font.size = Pt(28); tr.font.color.rgb = ACCENT
sp = doc.add_paragraph(); sr = sp.add_run(
    "Beleidsmonitor industrieel vastgoed — Today Development")
sr.font.size = Pt(13)
dp = doc.add_paragraph(); dr = dp.add_run("Routekaart voor doorontwikkeling")
dr.italic = True; dr.font.color.rgb = GRIJS

doc.add_paragraph()
alinea(doc, "Dit document beschrijft hoe de beleidsmonitor kan groeien. Het is "
            "opgebouwd uit vier assen (bronnen, zoekdiepte, geografie, techniek) "
            "plus drie thema's (diepere analyse, meldingen, overige functies) en "
            "een gefaseerde routekaart.")
alinea(doc, "Status bij schrijven: 2 bronnen, 3 gemeenten (Den Haag, Utrecht, "
            "Almere), ~1.300 signalen, online dashboard met wachtwoord.",
       italic=True, grijs=True)

# --- As 1 ---
kop(doc, "As 1 — Meer bronnen (breedte van informatie)")
alinea(doc, "Per bron: welke indicator(en) hij vult, of de data open is, de "
            "waarde en de bouwmoeite. Volgorde = aanbevolen prioriteit.")
tabel(doc, ["Bron", "Indicator(en)", "Open data?", "Waarde", "Moeite"],
      [
          ["Provinciale bekendmakingen (KOOP/SRU)", "1, 2, 3, 9", "Ja (al binnen bereik)", "Hoog", "Laag"],
          ["Netcongestie-capaciteitskaart (Netbeheer NL)", "4", "Ja (open kaart)", "Hoog", "Midden"],
          ["IBIS Bedrijventerreinen", "2", "Ja (via provincies)", "Hoog", "Midden"],
          ["Omgevingsplannen — beslissingen", "3, 7, 8", "Ja (al via bekendmakingen)", "Midden", "Laag"],
          ["Omgevingsplannen — live status (DSO)", "3, 7, 8", "Ja, maar complex (geo)", "Midden-hoog", "Hoog"],
          ["Kadaster — perceelgrenzen/adressen (PDOK/BAG)", "6", "Ja (gratis)", "Laag-midden", "Midden"],
          ["Kadaster — eigendom (BRK)", "6", "Betaald + privacy-beperkt", "Hoog", "Hoog/€"],
          ["Omgevingsdiensten — milieuvergunningen", "3, 10", "Wisselend", "Midden", "Hoog"],
          ["Gemeente-/provincienieuws (RSS)", "breed", "Ja", "Laag (ruis)", "Laag"],
      ],
      [5.0, 2.8, 3.4, 2.5, 2.3])
alinea(doc, "Nuances: de besluiten over omgevingsplannen komen al binnen via de "
            "bekendmakingen; alleen de live status per locatie (DSO) ontbreekt. "
            "Bij Kadaster zijn perceelgrenzen gratis, maar eigendom is betaald en "
            "privacy-beperkt.")
qw = doc.add_paragraph(); qw.add_run("Quick wins: ").bold = True
qw.add_run("provinciale bekendmakingen (lage moeite) en de netcongestie-kaart "
           "(vult indicator 4, nu het grootste gat).")

# --- As 2: zoekdiepte ---
kop(doc, "As 2 — Dieper zoeken (zoekdiepte)")
alinea(doc, "Hoe grondig de monitor zoekt, bepaalt of álle relevante signalen "
            "bovenkomen (recall). Een ladder van ondiep+goedkoop naar diep+duur:")
tabel(doc, ["Aanpak", "Recall", "Kosten", "Wat het is"],
      [
          ["1. Nu: trefwoorden + nieuwste N", "matig", "€", "Top 6-10 recente docs per zoekterm; mist de staart."],
          ["2. Dieper per zoekterm", "redelijk", "€€", "Alle treffers per term, niet alleen de nieuwste."],
          ["3. Bredere trefwoorden + synoniemen", "redelijk", "€", "Deels al gedaan (95 termen)."],
          ["4. Semantisch zoeken (embeddings)", "hoog", "€€", "Zoeken op betekenis; vangt docs zonder je trefwoorden."],
          ["5. Breed ophalen + AI-filter", "maximaal", "€€€", "Alle recente docs ophalen, AI filtert. Hoogste recall."],
          ["6. Volledige tekst i.p.v. fragment", "diepte/doc", "€€", "Lange documenten volledig laten lezen."],
          ["7. Verder terug in tijd", "historisch", "€€", "VANAF_JAAR verlagen → trends over jaren."],
          ["8. Tweede AI-pass (Opus)", "diepte oordeel", "€€", "Topdocumenten extra grondig analyseren."],
      ],
      [5.0, 2.2, 1.6, 7.2])
alinea(doc, "De grote drie keuzes: 4 (semantisch) en 5 (breed + AI-filter) "
            "verhogen écht de recall; 6/8 verdiepen de kwaliteit per document; "
            "2/3/7 zijn graduele stappen.")

# --- As 3: geografisch ---
kop(doc, "As 3 — Geografisch opschalen")
tabel(doc, ["Schaal", "Eenmalige kosten*", "Doorlooptijd**", "Techniek"],
      [
          ["3 gemeenten (nu)", "—", "—", "huidige opzet"],
          ["Regio (30-50 gemeenten)", "€200-400 (zuinig €50-100)", "uren", "Haiku + planning"],
          ["Provincie / Randstad", "€400-800", "~dag", "Batch API"],
          ["Heel NL (342)", "€1.000-2.000 (zuinig €300-600)", "dagen", "Batch API + cloud"],
      ],
      [4.6, 4.8, 2.6, 4.0])
alinea(doc, "* Ordegrootte, afhankelijk van model + documentaantal.  ** Vooral "
            "door API-throttling.", grijs=True)
adv = doc.add_paragraph(); adv.add_run("Advies: ").bold = True
adv.add_run("richt op de gemeenten waar Today actief is of wil zijn. Gericht "
            "opschalen vangt ~95% van de signalen tegen een fractie van de kosten.")

# --- As 4: techniek ---
kop(doc, "As 4 — Techniek die opschalen betaalbaar maakt")
bullets(doc, [
    "Goedkoper classificeren: claude-haiku-4-5 (~3× goedkoper), kortere tekst, prompt-caching.",
    "Batch API: ~50% goedkoper, gemaakt voor grote volumes.",
    "Automatische runs: wekelijks via Windows Taakplanner of GitHub Actions.",
    "Pipeline in de cloud: verzamelen hoeft niet op de eigen pc.",
])

# --- Diepere analyse ---
kop(doc, "Diepere analyse (wat je mét de data doet)")
bullets(doc, [
    "Trendanalyse over tijd (neemt transformatiedruk/netcongestie toe?).",
    "Vergelijking tussen gemeenten (benchmark vestigingsklimaat).",
    "Verandering-detectie (alleen tonen wat nieuw is sinds vorige run).",
    "Cross-document synthese (beleidsrichting per gemeente samenvatten).",
    "Koppeling aan Today's portefeuille (signaal bij eigen locatie weegt zwaarder).",
])

# --- Meldingen ---
kop(doc, "Meldingen")
alinea(doc, "Waarover wil je gewaarschuwd worden (de trigger):")
tabel(doc, ["Trigger", "Toelichting"],
      [
          ["Nieuw sterk risico (★4-5)", "Alleen de belangrijke risico's; weinig ruis."],
          ["Nieuw signaal in 'jouw' gemeente", "Bv. waar Today actief is."],
          ["Nieuw signaal op een indicator", "Bv. netcongestie (indicator 4)."],
          ["Op een trefwoord", "Opgeslagen zoekopdracht."],
          ["Statuswijziging", "Plan van ontwerp → vastgesteld."],
      ],
      [5.0, 11.0])
alinea(doc, "Hoe vaak: direct (kan druk) · wekelijkse/dagelijkse samenvatting "
            "(aanbevolen) · alleen bij iets belangrijks.")
tabel(doc, ["Kanaal", "Geschikt?", "Moeite"],
      [
          ["Microsoft Teams (webhook)", "Top voor een bedrijf (Microsoft 365)", "Laag"],
          ["E-mail", "Universeel", "Laag-midden"],
          ["In het dashboard ('Nieuw'-sectie)", "Geen extern kanaal", "Zeer laag / gratis"],
          ["Slack", "Als jullie Slack gebruiken", "Laag"],
      ],
      [6.0, 7.0, 3.0])

kop(doc, "Project-specifieke meldingen (aanbevolen — hoge waarde)", niveau=2)
alinea(doc, "Leg Today's projectlocaties vast en krijg een gerichte melding zodra "
            "er een signaal over die locatie binnenkomt — bv. een "
            "voorbereidingsbesluit (dat ontwikkeling kan bevriezen) in die buurt.")
bullets(doc, [
    "Lijst van projecten: gemeente + locatie-aanduidingen (naam, straat, wijk, kavel).",
    "Bij elk nieuw signaal: juiste gemeente én noemt het de projectlocatie?",
    "Zo ja → directe melding: '⚠️ Voorbereidingsbesluit nabij [project X] in [gemeente]'.",
])
tabel(doc, ["Niveau", "Hoe", "Haalbaarheid"],
      [
          ["A. Tekst-match", "projectnaam/straat/wijk in de signaaltekst herkennen", "Goed haalbaar"],
          ["B. Geo-match", "plangrenzen vergelijken met het perceel", "Geavanceerd (geo-data)"],
      ],
      [3.0, 9.0, 4.0])
alinea(doc, "Samenhang: meldingen vereisen dat de pipeline automatisch draait. De "
            "meldingen zelf zijn gratis; de kosten zitten in de geplande run.")

# --- Continu monitoren ---
kop(doc, "Continu monitoren (automatisch bijhouden)")
alinea(doc, "Nu draait de pipeline handmatig. Voor automatisch bijhouden moet de "
            "pipeline op een schema draaien én moeten de resultaten beschikbaar "
            "komen. Belangrijk: Streamlit Cloud kan de pipeline zelf niet draaien "
            "(het toont alleen het dashboard, en de opslag is tijdelijk); het "
            "verzamelen gebeurt elders, waarna de bijgewerkte database online komt.")
tabel(doc, ["Manier", "Pc uit = werkt?", "Kosten", "Moeite", "Past bij"],
      [
          ["Windows Taakplanner (eigen pc)", "Nee", "Gratis", "Laag", "snel starten, klein"],
          ["GitHub Actions (cloud)", "Ja", "Gratis (binnen limieten)", "Midden", "huidige GitHub + Streamlit"],
          ["Cloud-server (cron-job)", "Ja", "€ per maand", "Hoog", "bedrijfsbreed / groot"],
      ],
      [4.5, 2.7, 3.3, 1.8, 3.7])
alinea(doc, "Cadans: dagelijks · wekelijks (aanbevolen) · maandelijks. Wekelijks "
            "is meestal de juiste balans tussen actualiteit en kosten.")
bullets(doc, [
    ("Netcongestie", "altijd live — geen planning nodig."),
    ("Signalen (beleid)", "vereisen geplande pipeline-runs."),
    ("Meldingen", "hangen af van die geplande runs."),
])
alinea(doc, "Kosten per run: alleen nieuwe documenten worden geclassificeerd, dus "
            "een geplande run is goedkoop (paar cent tot ~€0,50, afhankelijk van "
            "volume en aantal gemeenten).")
ar = doc.add_paragraph(); ar.add_run("Aanbevolen opzet: ").bold = True
ar.add_run("GitHub Actions, wekelijks → de pipeline draait in de cloud, zet de "
           "bijgewerkte database terug in GitHub, Streamlit werkt zich vanzelf bij, "
           "en stuurt optioneel een melding bij nieuwe sterke risico's.")

# --- Overige functies ---
kop(doc, "Overige functies")
bullets(doc, [
    "Kaartweergave (signalen op een kaart van Nederland).",
    "Markeren & notities (afvinken, opmerkingen, status 'opgevolgd').",
    "Gebruikersrollen (wie ziet/doet wat bij meerdere gebruikers).",
    "Volledige-tekst zoeken (documenttekst meeopslaan).",
    "Deduplicatie over bronnen heen.",
    "Grafieken & Excel/PDF-export.",
    "Engelstalige versie.",
])

# --- Routekaart ---
kop(doc, "Routekaart (gefaseerd)")
bullets(doc, [
    ("Fase A — Verbreden & verdiepen", "provinciale bekendmakingen + netcongestie-kaart; eventueel zoekdiepte verhogen."),
    ("Fase B — Zuinig & schaalbaar maken", "Haiku + caching + Batch API."),
    ("Fase C — Gericht geografisch opschalen", "gemeentenlijst van Today's werkgebied → één grote run."),
    ("Fase D — Automatiseren & melden", "wekelijkse run + meldingen, inclusief project-specifieke alerts."),
])

# --- Kosten & privacy ---
kop(doc, "Kosten- en privacy-aandachtspunten")
bullets(doc, [
    "De bronnen zijn grotendeels gratis; kosten zitten in de AI-classificatie.",
    "Stel een maandlimiet in op console.anthropic.com (Billing → Limits).",
    "De API-sleutel hoort in secrets, nooit in de code of op GitHub.",
    "De online versie is een etalage; verzamelen gebeurt waar de pipeline draait.",
    "Kadaster-eigendom en sommige geo-data zijn betaald of privacy-beperkt.",
])

uit = "Uitbreidingsplan Beleidsmonitor Today Development.docx"
try:
    doc.save(uit)
    print("Document opgeslagen:", uit)
except PermissionError:
    uit2 = "Uitbreidingsplan Beleidsmonitor Today Development (bijgewerkt).docx"
    doc.save(uit2)
    print(f"Origineel stond open in Word; opgeslagen als: {uit2}")
